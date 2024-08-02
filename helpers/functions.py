import base64
import io
import random
from contextlib import suppress
from datetime import datetime
from typing import Union

from aiogram import Bot
from aiogram.client.session import aiohttp
from aiogram.enums import ContentType
from aiogram.exceptions import TelegramBadRequest
from aiogram.types import CallbackQuery, Message, InputMediaPhoto, InputMediaVideo, FSInputFile
from docx.shared import Inches
from openai import OpenAI

from data.context_vars import bot_session
from database import Order
from database.models import User
from utils.yookassa.api import payment
from docx import Document as DocxDocument



def russian_plural(n, forms):
    """
    Возвращает форму слова в зависимости от числа n.

    :param n: Число, в соответствии с которым выбирается форма
    :param forms: кортеж из трех форм слова (например, для слова 'яблоко': ('яблоко', 'яблока', 'яблок'))
    """
    if n % 10 == 1 and n % 100 != 11:
        return forms[0]
    elif 2 <= n % 10 <= 4 and (n % 100 < 10 or n % 100 >= 20):
        return forms[1]
    else:
        return forms[2]


async def edit_send_media(
        event: Union[CallbackQuery, Message],
        text: str,
        bot: Bot = bot_session.get('current_session'),
        reply_markup=None,
        media: dict = None,
        need_edit: bool = True
):
    bot: Bot = bot_session.get('current_session')
    message = event.message if isinstance(event, CallbackQuery) else event
    message_id = message.message_id
    chat_id = message.chat.id

    # Check if the message has media and if the sender is a bot
    have_media = message.video or message.photo or message.document
    is_bot = message.from_user.is_bot if message.from_user else True

    kwargs = {'reply_markup': reply_markup}

    if media is not None:
        send_method = message.answer_photo if media['type'] == ContentType.PHOTO else message.answer_video
        if have_media and is_bot:
            media_type = InputMediaPhoto if media['type'] == ContentType.PHOTO else InputMediaVideo
            with suppress(Exception):
                if need_edit:
                    return await bot.edit_message_media(
                        chat_id=chat_id,
                        message_id=message_id,
                        media=media_type(media=media['file_id'], caption=text),
                        reply_markup=reply_markup
                    )
                else:
                    return await send_method(
                        chat_id=chat_id,
                        media=media_type(media=media['file_id'], caption=text),
                        reply_markup=reply_markup
                    )
        else:
            if is_bot:
                await message.delete()
            try:
                return await send_method(media['file_id'], caption=text, **kwargs)
            except Exception:
                try:
                    file_extension = media['file_id'].split('.')[-1]
                    return await send_method(
                        media=FSInputFile(path=media['file_id'], filename=f"file.{file_extension}"),
                        caption=text, **kwargs
                    )
                except:
                    return await message.answer(text=text, **kwargs)

    else:
        kwargs['disable_web_page_preview'] = True
        if is_bot and not have_media:
            if need_edit:
                try:
                    return await message.edit_text(text=text, **kwargs)
                except TelegramBadRequest as error:
                    if error.message == 'Bad Request: message is not modified: specified new message content and reply markup are exactly the same as a current content and reply markup of the message':
                        return
                    else:
                        return await message.answer(text=text, **kwargs)
            else:
                return await message.answer(text=text, **kwargs)
        else:
            if have_media and is_bot:
                await message.delete()
            return await message.answer(text=text, **kwargs)


client = OpenAI(api_key="sk-proj-mxflsLySKuKrimvaPf3FT3BlbkFJyJ7KzjmCUbweZdJJm95a")


async def get_chat_response(prompt, information):
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system",
             "content": information},
            {"role": "user", "content": prompt},
        ]
    )
    content = response.choices[0].message.content
    return content


async def create_payment_link(user):
    payment_data = await payment(amount=1000, name=user.fio, phone=user.number)
    url = payment_data.confirmation.confirmation_url

    await Order(
        user=user.id,
        identy=payment_data.id,
        amount=1000
    ).insert()

    return url


def add_image_if_base64(doc, title, image_base64, width=Inches(3.0)):
    if image_base64:
        image_data = base64.b64decode(image_base64)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(image_data), width=width)
        paragraph.alignment = 1  # Center the image
        paragraph = doc.add_paragraph(title)
        paragraph.alignment = 1  # Center the title


def generate_user_report_in_memory(user: User):
    doc = DocxDocument()

    doc.add_heading(f'Досье на пользователя {user.full_name}', 0)

    doc.add_heading('Основная информация', level=1)
    doc.add_paragraph(f'Полное имя: {user.full_name}')
    doc.add_paragraph(f'Имя пользователя: {user.username}')
    doc.add_paragraph(f'Роль: {user.role}')
    doc.add_paragraph(f'Дата регистрации: {user.registration_date.strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(
        f'Последняя активность: {user.last_active.strftime("%Y-%m-%d %H:%M:%S") if user.last_active else "N/A"}')

    doc.add_heading('Верификация', level=1)
    doc.add_paragraph(f'Верификация на автомобиль: {"Да" if user.verification.verification_auto else "Нет"}')
    doc.add_paragraph(f'Верификация на документы: {"Да" if user.verification.verification_user else "Нет"}')

    doc.add_heading('Настройки', level=1)
    doc.add_paragraph(f'Язык: {user.settings.language}')

    doc.add_heading('Финансовая информация', level=1)
    doc.add_paragraph(f'Баланс: {user.balance} руб.')
    doc.add_paragraph(f'Подписка: {user.subscription if user.subscription else "Нет"}')

    doc.add_heading('Информация об автомобиле', level=1)
    if user.photo_auto_documents:
        doc.add_paragraph(f'Номера автомобиля: {user.photo_auto_documents.auto_number}')
        add_image_if_base64(doc, 'Перед автомобиля', user.photo_auto_documents.auto_front)
        add_image_if_base64(doc, 'Левая сторона автомобиля', user.photo_auto_documents.auto_left)
        add_image_if_base64(doc, 'Правая сторона автомобиля', user.photo_auto_documents.auto_right)
        add_image_if_base64(doc, 'Зад автомобиля', user.photo_auto_documents.auto_back)
        add_image_if_base64(doc, 'Салон спереди', user.photo_auto_documents.salon_front)
        add_image_if_base64(doc, 'Зад салона', user.photo_auto_documents.salon_back)

    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)

    return byte_stream


def generate_pincode(length=4):
    return ''.join([str(random.randint(0, 9)) for _ in range(length)])


async def make_tellcode_call(phone, pincode=None):
    if pincode is None:
        pincode = generate_pincode()

    url = "https://zvonok.com/manager/cabapi_external/api/v1/phones/tellcode/"
    payload = {
        'public_key': '55ea612b7d737fc2c35d2054e5da1fcb',
        'phone': phone,
        'campaign_id': 2008415037,
        'pincode': pincode
    }

    async with aiohttp.ClientSession() as session:
        async with session.post(url, data=payload) as response:
            print(response)
            if response.status == 200:
                response_data = await response.json()
                if response_data['status'] == "ok":
                    return response_data['data']['pincode']
                else:
                    return None
            else:
                return None